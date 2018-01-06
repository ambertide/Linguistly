import tkinter as tk
from tkinter import filedialog
import translation
from commonExpressions import CommonExpressions, prepare, save, strip, draw
from docx import Document


class Latte:
    def __init__(self, master):
        self.master = master
        self.language = translation.main.langCode()
        master.title("LinguaTK")
        master.bind("a", self.advanced)
        master.bind("A", self.advanced)
        self.file_number_var = 1
        self.file_directory = ""
        self.file_directory_label = tk.Label(master,
            text = translation.Cover.pleaseSelect[self.language])
        self.file_directory_button = tk.Button(master,
            text = translation.Cover.browse[self.language], command = lambda: self.browse())
        self.selectorVar = tk.StringVar(master)
        self.selectorVar.set("csv")
        self.advencedVar = tk.IntVar(master)
        self.checkBoxVar = tk.IntVar(master)
        self.output_method_select = tk.OptionMenu(master, self.selectorVar,
            "txt", "sqlite3", "xlsx", "csv")
        self.addButton = tk.Button(master, text= "+", command= lambda: self.add())
        self.is_turkish = tk.Checkbutton(master, text= translation.Cover.isTurkish[self.language],
            variable = self.checkBoxVar)
        self.output_finalize_button = tk.Button(master, text=translation.Cover.analyise[self.language],
            state = tk.DISABLED, command= lambda: self.commence())
        self.output_draw_button = tk.Button(master, text=translation.Cover.plusDraw[self.language],
            state= tk.DISABLED, command= lambda: self.commence(1))
        self.file_directory_label.grid(row=0, column=0, columnspan=2,
            sticky=tk.W + tk.E + tk.S + tk.N)
        self.file_directory_button.grid(row=0, column=2,
            sticky=tk.W + tk.E + tk.S + tk.N)
        self.addButton.grid(row=0, column=3, sticky=tk.W + tk.E + tk.S + tk.N)
        self.output_method_select.grid(row=1, column=0,
            sticky=tk.W + tk.E + tk.S + tk.N)
        self.is_turkish.grid(row=1, column=1)
        self.suffix_clean = tk.IntVar(master)
        self.suffix_checkBox = tk.Checkbutton(master, text= translation.Cover.isSuffix[self.language],
            variable = self.suffix_clean)
        file_ = open("settings.lingua", "r")
        self.custom_remove = file_.read()
        file_.close()
        self.suffix_checkBox.grid(row=1, column=2, columnspan=2)
        self.output_finalize_button.grid(row=2, column=0, columnspan=3,
            sticky=tk.W + tk.E + tk.S + tk.N)
        self.output_draw_button.grid(row=2, column=3,
            sticky=tk.W + tk.E + tk.S + tk.N)
        self.custom_words = []
    def browse(self):
        self.file_number_var = 1
        options = {}
        options['defaultextension'] = '.txt'
        options['filetypes'] = [(translation.Browse.textFiles[self.language], ".txt"),
                                (translation.Browse.docFiles[self.language], ".docx")]
        options['title'] = translation.Browse.selectFiles[self.language]
        self.file_directory = filedialog.askopenfilename(**options)
        self.file_directory_label['text'] = self.file_directory
        self.output_finalize_button['state'] = 'normal'
        self.output_draw_button['state'] = 'normal'

    def add(self):
        addScreen = tk.Toplevel()
        addScreen.title(translation.AddFiles.addFile[self.language])
        self.file_number_var = 0
        self.filenames = tk.Entry(addScreen)
        self.filenames.grid(row=0, column=0, columnspan=4, sticky=tk.W + tk.E + tk.S + tk.N)
        self.filenames.insert(tk.END, translation.AddFiles.info[self.language])
        self.addFiles = tk.Button(addScreen, text=translation.AddFiles.add[self.language],
            command= lambda: self.add_command())
        self.addFiles.grid(row=1, column=0, columnspan=4, sticky=tk.W + tk.E + tk.S + tk.N)

    def add_command(self):
        self.file_directory = self.filenames.get()
        self.file_directory_label['text'] = self.file_directory
        self.output_finalize_button['state'] = 'normal'
        self.output_draw_button['state'] = 'normal'

    def advanced(self, event=None):
        advencedOp = tk.Toplevel()
        conjunction = tk.Checkbutton(advencedOp, text=translation.Conjunctions.removeConjunctions[self.language],
            variable= self.advencedVar)
        conjunction.grid(row=0, column=0)
        special = tk.Entry(advencedOp)
        file_ = open("settings.lingua", "r")
        settings = file_.read()
        file_.close()
        special.insert(tk.END, settings)
        special.grid(row=1, column=0)
        save = tk.Button(advencedOp, text=translation.Conjunctions.addWords[self.language], command= lambda: self.savewords(special.get()))
        save.grid(row=2, column=0)

    def savewords(self, input_):
        file_ = open("settings.lingua", "w")
        file_.write(input_)
        file_.close()
        self.custom_words = input_.split(",")
    def commence(self, is_draw = 0):
        if self.file_number_var == 1:
            extension = self.file_directory.split(".")
            initdata = ""
            findata = []
            if extension[1] == "txt":
                file = open(self.file_directory, "r")
                initdata = file.read()
                file.close()
            elif extension[1] == "docx":
                initdata = ""
                file = Document(self.file_directory)
                parags = file.paragraphs
                for i in range(len(parags)):
                    initdata += parags[i].text
            if self.suffix_clean.get() == 1:
                if self.checkBoxVar.get() == 1:
                    initdata = strip(initdata, "Turkish", "suffix")
                elif self.checkBoxVar.get() == 0:
                    initdata = strip(initdata, "English", "suffix")
            if self.advencedVar.get() == 1:
                if self.checkBoxVar.get() == 1:
                    initdata = strip(initdata, "Turkish", "conjunctions")
                elif self.checkBoxVar.get() == 0:
                    initdata = strip(initdata, "English", "conjunctions")
            if self.custom_words != []:
                initdata = strip(initdata, toBeRemoved = self.custom_words)
            if self.checkBoxVar.get() == 1:
                findata = prepare(initdata, "yes")
            elif self.checkBoxVar.get() == 0:
                findata = prepare(initdata)
            indexed = {}
            indata = list(set(findata))
            for i in range(len(indata)):
                indexed[indata[i]] = findata.count(indata[i])
            keysindex = list(indexed)
            save(indexed, keysindex,
                outputtype = self.selectorVar.get())
            if is_draw == 1:
                draw(indexed)
        elif self.file_number_var == 0:
            fileNames = self.file_directory.split(",")
            for i in range(len(fileNames)):
                self.file_directory = fileNames[i]
                extension = self.file_directory.split(".")
                initdata = ""
                findata = []
                if extension[1] == "txt":
                    file = open(self.file_directory, "r")
                    initdata = file.read()
                    file.close()
                elif extension[1] == "docx":
                    initdata = ""
                    file = Document(self.file_directory)
                    parags = file.paragraphs
                    for i in range(len(parags)):
                        initdata += parags[i].text
                if self.suffix_clean.get() == 1:
                    if self.checkBoxVar == 1:
                        initdata = strip(initdata, "Turkish", "suffix")
                    elif self.checkBoxVar == 0:
                        initdata = strip(initdata, "English", "suffix")
                if self.advencedVar.get() == 1:
                    if self.checkBoxVar.get() == 1:
                        initdata = strip(initdata, "Turkish", "conjunctions")
                    elif self.checkBoxVar.get() == 0:
                        initdata = strip(initdata, "English", "conjunctions")
                if self.custom_words != []:
                    initdata = strip(initdata, toBeRemoved = self.custom_words)
                if self.checkBoxVar.get() == 1:
                    findata = prepare(initdata, "yes")
                elif self.checkBoxVar.get() == 0:
                    findata = prepare(initdata)
                indexed = {}
                indata = list(set(findata))
                for i in range(len(indata)):
                    indexed[indata[i]] = findata.count(indata[i])
                keysindex = list(indexed)
                save(indexed, keysindex,
                    outputtype = self.selectorVar.get())
                if is_draw == 1:
                    draw(indexed)
root = tk.Tk()
cappucino = Latte(root)
root.mainloop()
